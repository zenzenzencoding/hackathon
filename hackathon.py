#!/usr/bin/env python
# encoding: utf-8

"""
@version: 0.1
@author: wang.zheng@Ctrip.com
@license: Apache Licence 
@file: hackathon.py
@time: 2017/3/23 17:25
"""
import win32com
import docx
from win32com.client import Dispatch,constants
from docx import Document
import os
import subprocess
def parse_docx(filename):
    """读取docx
    """
    document = Document(filename)
    tables = document.tables
    zws=[]
    bgs=[]
    for zw in document.paragraphs:
        zws.append(zw)
    for table in tables:
        sum = len(table._cells)#总表格数目
        columns = table._column_count #列数
        rows = sum / columns #列数
        tabletext =[]
        for row in range(rows):
            for column in range(columns):
                tabletext.append(table.cell(row,column).text)
        bgs.append(tabletext)

def parse_doc(filename):
    '''
    解析doc
    '''
    wordapp = win32com.client.Dispatch('Word.Application')
    doc = wordapp.Documents.Open(filename)
    docx=wordapp.Documents
    tables = doc.Tables # 根据文件中的图表选择信息
    for table in tables:
        docx.Add(table)
    docx.SaveAs("1.docx")
        # rows=len(table.Rows)
        # cols =len(table.Rows[0].Cells)
        # for row in range(rows):
        #     for col in range(cols):
        #         print table.Rows[row].Cells[col].Range.Text

if __name__ == '__main__':
    w = win32com.client.Dispatch('Word.Application')
    # 遍历文件
    PATH = r"E:\zenwan\Desktop\gysbj" # windows文件路径
    doc_files = os.listdir(PATH)
    for doc in doc_files:
        fn = PATH+"\\"+doc.decode("gbk").encode("utf-8")
        if os.path.splitext(doc)[1] == '.docx':
            #parse_docx(fn)
            pass
        if os.path.splitext(doc)[1] == '.doc':
            parse_doc(fn)

